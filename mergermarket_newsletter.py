#!/usr/bin/env python3
"""
Mergermarket Daily Newsletter Automation

Downloads the Mergermarket intelligence report, converts it to a formatted
Word document, and opens a pre-composed Outlook email for manual review.

Usage:
    python mergermarket_newsletter.py                      # full pipeline
    python mergermarket_newsletter.py --dry-run FILE.xlsx  # skip download
    python mergermarket_newsletter.py --date 2024-01-08    # override date
"""

from __future__ import annotations

import argparse
import ctypes
import logging
import os
import sys
import tempfile
import time
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
# Paths  (no hardcoded user or drive paths — works on any Windows machine)
# ---------------------------------------------------------------------------

# Temporary files: raw Excel download, debug screenshots/JSON
TEMP_DIR = Path(tempfile.gettempdir()) / "mergermarket"

# Persistent output: Word document, log file
OUTPUT_DIR = Path.home() / "Downloads"
LOG_FILE = OUTPUT_DIR / "mergermarket_log.txt"
SEPARATOR = "--------------------------"
MERGERMARKET_URL = "https://www.mergermarket.com/intelligence/intelligence.asp"
GEOGRAPHIES = ["Austria", "Germany", "Switzerland"]
EMAIL_RECIPIENT = "CASE_Germany"
EMAIL_INTRO = (
    "Guten Morgen,\n\n"
    "anbei ein aktueller Auszug aus Mergermarket.\n\n"
    "Mergermarket:\n"  # rendered bold; blank line added automatically after
)
FRIDAY_INTRO = (
    "Guten Morgen,\n\n"
    "anbei die laufenden Verfahren des Bundeskartellamts und ein aktueller "
    "Auszug aus Mergermarket.\n"
)

BKA_URL = (
    "https://www.bundeskartellamt.de/SiteGlobals/Forms/Suche/"
    "LaufendeVerfahren/LaufendeVerfahren_Formular.html"
)

# Boilerplate text patterns to skip when parsing the Excel report
BOILERPLATE_PREFIXES = ("handelsblatt", "mergermarket", "---", "===")

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------


def setup_logging() -> logging.Logger:
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    logger = logging.getLogger("mergermarket")
    logger.setLevel(logging.DEBUG)
    fmt = logging.Formatter(
        "%(asctime)s [%(levelname)s] %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )
    fh = logging.FileHandler(LOG_FILE, encoding="utf-8")
    fh.setFormatter(fmt)
    logger.addHandler(fh)
    ch = logging.StreamHandler()
    ch.setFormatter(fmt)
    logger.addHandler(ch)
    return logger


log = setup_logging()


def show_error(title: str, message: str) -> None:
    """Show a Windows MessageBox and log the error."""
    log.error(f"{title}: {message}")
    try:
        ctypes.windll.user32.MessageBoxW(0, message, title, 0x10)  # MB_ICONERROR
    except Exception:
        pass  # Non-Windows environment


# ---------------------------------------------------------------------------
# Date utilities
# ---------------------------------------------------------------------------


def get_run_date(override: Optional[str] = None) -> date:
    if override:
        return datetime.strptime(override, "%Y-%m-%d").date()
    return date.today()


def get_date_range(run_date: date) -> tuple[Optional[date], Optional[date]]:
    """Return (date_from, date_to) for Monday; (None, None) for Tue–Fri."""
    if run_date.weekday() == 0:  # Monday
        return run_date - timedelta(days=3), run_date  # Friday → Monday
    return None, None


def fmt_dmy(d: date) -> str:
    """Format as dd/mm/yyyy for Mergermarket date fields."""
    return d.strftime("%d/%m/%Y")


# ---------------------------------------------------------------------------
# Step 1 – Browser Automation (Playwright)
# ---------------------------------------------------------------------------


# ── Diagnostic helpers ───────────────────────────────────────────────────────

def _dump_page_state(page, label: str) -> None:
    """
    Save a full-page screenshot and a JSON DOM dump to C:\\Temp\\mm_debug_<label>.*

    Called automatically before every key interaction so that when a selector
    fails, we have complete visibility into what was actually on the page.
    Logs a human-readable summary of all selects / buttons / inputs / iframes.
    """
    import json as _json

    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    base = TEMP_DIR / f"mm_debug_{label}"

    # Screenshot (always from the top-level page object)
    try:
        page.screenshot(path=str(base.with_suffix(".png")), full_page=True)
    except Exception as exc:
        log.debug(f"Screenshot failed ({label}): {exc}")

    # DOM dump (evaluate in whichever context was passed — page or frame)
    try:
        data = page.evaluate("""() => {
            const selects = Array.from(document.querySelectorAll('select')).map(s => ({
                name: s.name, id: s.id, multiple: s.multiple,
                options: Array.from(s.options).slice(0, 30).map(o => o.text.trim())
            }));
            const buttons = Array.from(document.querySelectorAll(
                'button, input[type="button"], input[type="submit"]'
            )).map(b => ({
                tag: b.tagName.toLowerCase(),
                type: b.type || '',
                id: b.id,
                name: b.name || '',
                value: (b.value || '').slice(0, 80),
                text: (b.textContent || '').trim().slice(0, 80)
            }));
            const inputs = Array.from(document.querySelectorAll('input')).map(i => ({
                type: i.type, name: i.name || '', id: i.id,
                value: (i.value || '').slice(0, 60),
                placeholder: i.placeholder || ''
            }));
            const iframes = Array.from(document.querySelectorAll('iframe')).map(f => ({
                src: f.src, id: f.id, name: f.name
            }));
            return {
                url: location.href, title: document.title,
                selects, buttons, inputs, iframes
            };
        }""")

        with open(str(base.with_suffix(".json")), "w", encoding="utf-8") as fh:
            _json.dump(data, fh, indent=2, ensure_ascii=False)

        log.info(
            f"[DIAG {label}] {data['url'][:80]}  "
            f"selects={len(data['selects'])}  "
            f"buttons={len(data['buttons'])}  "
            f"inputs={len(data['inputs'])}  "
            f"iframes={len(data['iframes'])}"
        )
        log.info(f"  → screenshot: {base}.png")
        log.info(f"  → DOM dump  : {base}.json")

        # Print every select's options so selector issues are immediately visible
        for s in data["selects"]:
            log.info(f"  <select> name={s['name']!r} id={s['id']!r} "
                     f"multiple={s['multiple']} options={s['options'][:6]}")

        # Print every button
        for b in data["buttons"]:
            log.info(f"  <{b['tag']}> type={b['type']!r} "
                     f"value={b['value']!r} text={b['text']!r} "
                     f"id={b['id']!r} name={b['name']!r}")

        # Warn about iframes — selectors won't work across frame boundaries
        for f in data["iframes"]:
            log.warning(f"  !! iframe detected: id={f['id']!r} src={f['src']!r}")

    except Exception as exc:
        log.debug(f"DOM dump failed ({label}): {exc}")


def _find_form_context(page):
    """
    Return the Playwright context (Page or Frame) that contains the search form.

    If the form lives inside an <iframe> — common in legacy enterprise apps —
    selectors run against the main page document won't find anything.
    This function checks every frame for <select> elements and returns the
    first frame that has them.  Falls back to the main page.
    """
    # Check main document first
    try:
        n = page.evaluate("() => document.querySelectorAll('select').length")
        if n > 0:
            log.info(f"Form context: main page ({n} selects found)")
            return page
    except Exception:
        pass

    # Walk child frames
    for frame in page.frames[1:]:
        try:
            n = frame.evaluate("() => document.querySelectorAll('select').length")
            if n > 0:
                log.info(f"Form context: iframe '{frame.url}' ({n} selects found)")
                return frame
        except Exception:
            continue

    log.warning(
        "Form context: no <select> elements found anywhere on the page.\n"
        "The page may still be loading, require a login, or use a non-standard layout.\n"
        "Check the screenshot in C:\\Temp\\mm_debug_01_after_login.png"
    )
    return page  # best-effort fallback


# ── Main download orchestration ──────────────────────────────────────────────

def download_mergermarket_report(
    run_date: date,
    output_path: Path,
    *,
    headless: bool = False,
    is_friday: bool = False,
) -> tuple[Path, list[dict]]:
    """
    Navigate Mergermarket via Playwright, configure the search, and download
    the Unformatted Report Excel file to *output_path*.

    Diagnostic screenshots + DOM dumps are saved to C:\\Temp\\mm_debug_*.{png,json}
    at every key step so selector problems are immediately diagnosable.
    """
    try:
        from playwright.sync_api import TimeoutError as PWTimeout, sync_playwright
    except ImportError:
        show_error(
            "Mergermarket – Missing dependency",
            "Playwright is not installed.\n\nRun:\n  pip install playwright\n  playwright install chromium",
        )
        sys.exit(1)

    mm_user = os.environ.get("MM_USERNAME", "")
    mm_pass = os.environ.get("MM_PASSWORD", "")
    date_from, date_to = get_date_range(run_date)

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=headless)
        ctx_browser = browser.new_context(accept_downloads=True)
        page = ctx_browser.new_page()

        # Navigate directly to the intelligence search page.
        # If the session is not authenticated, the server redirects automatically
        # to id.ionanalytics.com/signin — _handle_login deals with that and
        # then navigates back here.  Using domcontentloaded (not networkidle)
        # so the goto does not hang while the SSO redirect chain is in flight.
        log.info(f"Navigating to: {MERGERMARKET_URL}")
        try:
            page.goto(MERGERMARKET_URL, wait_until="domcontentloaded", timeout=30_000)
        except PWTimeout:
            _dump_page_state(page, "00_timeout")
            show_error("Mergermarket – Timeout", "Could not load the Mergermarket page within 30 s.")
            browser.close()
            raise

        _handle_login(page, mm_user, mm_pass)

        # Wait for the search form, then take a diagnostic snapshot
        try:
            page.wait_for_selector("form", timeout=15_000)
        except Exception:
            pass
        page.wait_for_timeout(1_000)  # allow any late JS rendering to settle
        _dump_page_state(page, "01_after_login")

        # Determine whether the form lives in the main document or an iframe
        ctx = _find_form_context(page)

        # ── Date range ───────────────────────────────────────────────────────
        if date_from and date_to:
            log.info(f"Setting date range: {fmt_dmy(date_from)} – {fmt_dmy(date_to)}")
            _set_date_range(ctx, date_from, date_to)
        else:
            log.info("Selecting 'Last 24 Hours' …")
            _select_last_24h(ctx)

        # ── Geography ────────────────────────────────────────────────────────
        log.info("Selecting geographies: Austria, Germany, Switzerland …")
        _dump_page_state(page, "02_before_geography")
        _select_geographies(ctx, GEOGRAPHIES)

        # ── Search (use the last Search button — after the Geography section) ─
        log.info("Submitting search …")
        _dump_page_state(page, "03_before_search")
        ctx.evaluate("window.scrollTo(0, document.body.scrollHeight)")
        ctx.wait_for_timeout(300)

        search_btns = ctx.locator("input[value='Search']")
        if search_btns.count() > 0:
            search_btns.last.scroll_into_view_if_needed()
            search_btns.last.click()
            log.info(f"Clicked Search button ({search_btns.count()} found, used last)")
        else:
            clicked = _try_click(ctx, [
                "button:text-is('Search')",
                "input[type='submit']",
            ])
            if not clicked:
                _dump_page_state(page, "03b_search_not_found")
                raise RuntimeError(
                    "Search button not found. "
                    f"See C:\\Temp\\mm_debug_03b_search_not_found.png"
                )

        page.wait_for_load_state("networkidle", timeout=30_000)

        # ── Download ─────────────────────────────────────────────────────────
        log.info("Initiating report download …")
        _dump_page_state(page, "04_results_page")
        downloaded = _trigger_download(page, ctx, output_path)
        log.info(f"Report saved to: {downloaded}")

        bka_data: list[dict] = []
        if is_friday:
            log.info("Friday run — scraping Bundeskartellamt …")
            bka_data = scrape_bundeskartellamt(page)

        browser.close()

    return downloaded, bka_data


# ── Form-interaction helpers (accept Page or Frame as `ctx`) ─────────────────

def _handle_login(page, username: str, password: str) -> None:
    """
    Handle ION Analytics SSO login for Mergermarket.

    Detection: URL-based, not DOM-based.  The SSO provider redirects to
    id.ionanalytics.com/signin?onSuccess=… before the actual Mergermarket
    page loads.  A password-field check would miss this because the field
    is injected dynamically after the username step.

    Two-step flow:
      1. Fill username/email  →  click Continue / Next / Sign in
      2. Wait for password field to appear  →  fill  →  click Sign in
      3. Wait for redirect back to *.mergermarket.com/*
    """
    url = page.url.lower()

    # Detect login page by URL (primary signal)
    on_login_page = (
        "ionanalytics.com" in url
        or "/signin" in url
        or "/login" in url
        or "/auth/" in url
        or "mergermarket.com" not in url   # still on a redirect/SSO host
    )

    if not on_login_page:
        log.info(f"Already on Mergermarket — skipping login ({page.url[:70]})")
        return

    log.info(f"Login / SSO page detected: {page.url[:80]}")

    if not username or not password:
        show_error(
            "Mergermarket – Login Required",
            "Redirected to login page but MM_USERNAME / MM_PASSWORD are not set.\n"
            "Set them as environment variables:\n"
            "  setx MM_USERNAME \"you@company.com\"\n"
            "  setx MM_PASSWORD \"yourpassword\"",
        )
        raise RuntimeError("MM_USERNAME / MM_PASSWORD not configured")

    # ── Step 1: username / email ─────────────────────────────────────────────
    for sel in [
        "input[type='email']",
        "input[name='email']",
        "input[name='username']",
        "input[name='identifier']",
        "input[type='text']",
    ]:
        if page.query_selector(sel):
            page.fill(sel, username)
            log.info(f"Filled username/email via {sel!r}")
            break
    else:
        log.warning("Could not find a username/email field on the login page")

    # Click Continue / Next — many SSO flows show the password on the next screen
    _try_click(page, [
        "button:has-text('Continue')",
        "button:has-text('Next')",
        "input[value='Continue']",
        "input[value='Next']",
        "button[type='submit']",
        "input[type='submit']",
    ])

    # Wait up to 5 s for the password field to appear (two-step flow)
    try:
        page.wait_for_selector("input[type='password']", timeout=5_000)
        log.info("Password field appeared after Continue step")
    except Exception:
        pass  # Field may already be visible (single-step form)

    # ── Step 2: password ─────────────────────────────────────────────────────
    pw_field = page.query_selector("input[type='password']")
    if pw_field:
        page.fill("input[type='password']", password)
        log.info("Filled password field")
    else:
        log.warning("Password field still not visible — login may fail")

    # Submit: try CSS selectors first, then JS scan (case-insensitive), then Enter key.
    # The Enter-key fallback is the most reliable across all SSO implementations.
    submitted = _try_click(page, [
        "button[type='submit']",
        "input[type='submit']",
        "button:has-text('Sign in')",
        "button:has-text('Sign In')",
        "button:has-text('Log in')",
        "button:has-text('Login')",
    ])

    if not submitted:
        # JS fallback: find any submit/button whose text contains a login keyword
        submitted = page.evaluate("""() => {
            const terms = ['sign in', 'log in', 'login', 'submit', 'continue'];
            // Prefer explicit submit buttons
            const submits = Array.from(
                document.querySelectorAll('button[type="submit"], input[type="submit"]')
            );
            if (submits.length) { submits[submits.length - 1].click(); return true; }
            // Fall back to any button with matching text
            for (const el of document.querySelectorAll('button, input[type="button"]')) {
                const t = (el.textContent || el.value || '').trim().toLowerCase();
                if (terms.some(kw => t.includes(kw))) { el.click(); return true; }
            }
            return false;
        }""")

    if not submitted and pw_field:
        # Last resort: press Enter in the password field
        log.info("No submit button found — pressing Enter in password field")
        pw_field.press("Enter")

    log.info("Login submitted — waiting 3 s then navigating to intelligence page …")

    # ── Navigate immediately after submit — no URL polling ───────────────────
    # The SSO redirect can take an unpredictable amount of time.  Rather than
    # waiting for a specific URL pattern (which can block for 30 s+), we
    # simply give the auth flow 3 seconds to set its cookies, then navigate
    # directly to the target page.  If the session was not actually established,
    # the page will redirect back to the SSO and the login block will re-run.
    page.wait_for_timeout(3_000)
    log.info(f"Navigating to: {MERGERMARKET_URL}")
    page.goto(MERGERMARKET_URL, wait_until="domcontentloaded", timeout=30_000)
    log.info(f"Ready: {page.url[:80]}")


def _set_date_range(ctx, date_from: date, date_to: date) -> None:
    """
    Activate the 'Date From / Date To' radio and fill in both text inputs.

    Page structure (from live screenshot):
      ● [radio]  Last 24 Hours  [<select>]
      ○ [radio]  Date From [text input]  Date To [text input]  Clear Date

    The custom-date radio is the one whose siblings do NOT include a <select>.
    """
    ctx.evaluate("""() => {
        const radios = Array.from(document.querySelectorAll('input[type="radio"]'));
        for (const r of radios) {
            let el = r.nextElementSibling;
            let hasSelect = false;
            for (let i = 0; i < 4; i++) {
                if (!el) break;
                if (el.tagName === 'SELECT') { hasSelect = true; break; }
                el = el.nextElementSibling;
            }
            if (!hasSelect) { r.click(); return; }
        }
    }""")
    ctx.wait_for_timeout(400)

    ctx.evaluate(f"""() => {{
        const inputs = Array.from(document.querySelectorAll('input[type="text"]'));
        let fromEl = null, toEl = null;
        for (const inp of inputs) {{
            const context = (inp.parentElement?.textContent ?? '') +
                            (inp.parentElement?.parentElement?.textContent ?? '');
            if (!fromEl && context.includes('Date From')) {{ fromEl = inp; continue; }}
            if (!toEl   && context.includes('Date To'))   {{ toEl   = inp; continue; }}
        }}
        if (!fromEl && inputs[0]) fromEl = inputs[0];
        if (!toEl   && inputs[1]) toEl   = inputs[1];

        const set = (el, val) => {{
            if (!el) return;
            el.value = val;
            ['input', 'change', 'blur'].forEach(ev =>
                el.dispatchEvent(new Event(ev, {{bubbles: true}})));
        }};
        set(fromEl, '{fmt_dmy(date_from)}');
        set(toEl,   '{fmt_dmy(date_to)}');
    }}""")
    log.info(f"Date range set: {fmt_dmy(date_from)} → {fmt_dmy(date_to)}")


def _select_last_24h(ctx) -> None:
    """
    Select 'Last 24 Hours' from the daterange <select> dropdown.

    The page has a radio + <select> pair for the "quick range" option.
    Clicking the radio is not enough — the <select> retains its previous
    value (e.g. 'Last 12 Months').  We find the <select> that contains a
    'Last 24 Hours' option and set it explicitly, then also click the
    associated radio so the form registers the correct mode.
    """
    result = ctx.evaluate("""() => {
        // Find the <select> whose options include 'Last 24 Hours'
        const sel = Array.from(document.querySelectorAll('select')).find(s =>
            Array.from(s.options).some(o => o.text.trim() === 'Last 24 Hours')
        );
        if (!sel) return false;

        // Set the value to the 'Last 24 Hours' option
        const opt = Array.from(sel.options).find(o => o.text.trim() === 'Last 24 Hours');
        sel.value = opt.value;
        sel.dispatchEvent(new Event('change', {bubbles: true}));

        // Also click the radio button that owns this <select>, if present
        let el = sel.previousElementSibling;
        for (let i = 0; i < 4; i++) {
            if (!el) break;
            if (el.tagName === 'INPUT' && el.type === 'radio') { el.click(); break; }
            el = el.previousElementSibling;
        }
        return opt.value;
    }""")

    if result is False:
        log.warning("_select_last_24h: <select> with 'Last 24 Hours' option not found — check page structure")
    else:
        log.info(f"Date mode set to 'Last 24 Hours' (option value={result!r})")
    ctx.wait_for_timeout(300)


def _select_geographies(ctx, countries: list[str]) -> None:
    """
    Geography uses 4 cascading <select> elements:
      Col 1: Continent → Col 2: Sub-region → Col 3: Country

    We trigger each level by setting option.selected and dispatching 'change'.
    """
    # Step 1 – select 'Europe' in the continent column
    ok1 = ctx.evaluate("""() => {
        const sels = Array.from(document.querySelectorAll('select'));
        const s = sels.find(sel => {
            const opts = Array.from(sel.options).map(o => o.text.trim());
            return opts.includes('Europe') && opts.includes('Americas');
        });
        if (!s) return false;
        Array.from(s.options).forEach(o => { o.selected = o.text.trim() === 'Europe'; });
        s.dispatchEvent(new Event('change', {bubbles: true}));
        return true;
    }""")
    if not ok1:
        log.warning(
            "Geography: continent <select> not found. "
            "Check mm_debug_02_before_geography.json — the select with options "
            "[Africa, Americas, Asia, Europe, Middle East] was not present."
        )
        return
    log.info("Geography: continent 'Europe' selected")

    try:
        ctx.wait_for_function(
            "() => Array.from(document.querySelectorAll('select option'))"
            "       .some(o => o.text.trim() === 'Western Europe')",
            timeout=6_000,
        )
    except Exception:
        ctx.wait_for_timeout(2_000)

    # Step 2 – select 'Western Europe'
    ok2 = ctx.evaluate("""() => {
        const sels = Array.from(document.querySelectorAll('select'));
        const s = sels.find(sel =>
            Array.from(sel.options).some(o => o.text.trim() === 'Western Europe')
        );
        if (!s) return false;
        Array.from(s.options).forEach(o => { o.selected = o.text.trim() === 'Western Europe'; });
        s.dispatchEvent(new Event('change', {bubbles: true}));
        return true;
    }""")
    if not ok2:
        log.warning("Geography: 'Western Europe' not found after cascading from Europe")
        return
    log.info("Geography: sub-region 'Western Europe' selected")

    try:
        ctx.wait_for_function(
            "() => Array.from(document.querySelectorAll('select option'))"
            "       .some(o => o.text.trim() === 'Germany')",
            timeout=6_000,
        )
    except Exception:
        ctx.wait_for_timeout(2_000)

    # Step 3 – multi-select target countries
    result = ctx.evaluate("""(targets) => {
        const sels = Array.from(document.querySelectorAll('select'));
        const s = sels.find(sel =>
            targets.some(c => Array.from(sel.options).some(o => o.text.trim() === c))
        );
        if (!s) return {found: false};
        const matched = [];
        Array.from(s.options).forEach(o => {
            if (targets.includes(o.text.trim())) { o.selected = true; matched.push(o.text.trim()); }
        });
        s.dispatchEvent(new Event('change', {bubbles: true}));
        return {found: true, matched};
    }""", countries)

    if not result.get("found"):
        log.warning("Geography: country <select> not found — Austria/Germany/Switzerland absent")
    else:
        matched = result.get("matched", [])
        missing = [c for c in countries if c not in matched]
        log.info(f"Geography: selected {matched}")
        if missing:
            log.warning(f"Geography: options not found for {missing}")


def _try_click(ctx, selectors: list[str]) -> bool:
    """Try each selector in order; click the first one found. Returns True on success."""
    for sel in selectors:
        try:
            elem = ctx.query_selector(sel)
            if elem:
                elem.click()
                return True
        except Exception:
            continue
    log.warning(f"None of the click selectors matched: {selectors}")
    return False


def _js_click_by_text(ctx, *phrases: str) -> bool:
    """
    Case-insensitive JS fallback: click the first <a>, <button>, or <input>
    whose visible text / value contains any of the given phrases.
    Returns True if something was clicked.
    """
    return ctx.evaluate("""(phrases) => {
        const lower = phrases.map(p => p.toLowerCase());
        const els = Array.from(document.querySelectorAll(
            'a, button, input[type="button"], input[type="submit"]'
        ));
        for (const el of els) {
            const t = (el.textContent || el.value || '').trim().toLowerCase();
            if (lower.some(p => t.includes(p))) { el.click(); return true; }
        }
        return false;
    }""", list(phrases))


def _validate_excel_download(path: Path) -> str:
    """
    Verify the downloaded file is a real Excel file by inspecting magic bytes.

    Accepts:
      - b'PK'                → .xlsx  (ZIP-based Office Open XML)
      - b'\\xd0\\xcf\\x11\\xe0' → .xls   (OLE2 / BIFF, Office 97–2003)

    Returns 'xlsx' or 'xls'.  Raises RuntimeError with a content preview if
    the file is neither (e.g. an HTML error page returned by the server).
    """
    try:
        with open(path, "rb") as fh:
            header = fh.read(4)
    except OSError as exc:
        raise RuntimeError(f"Download fehlgeschlagen – Datei nicht lesbar: {exc}") from exc

    if header[:2] == b"PK":
        return "xlsx"
    if header == b"\xd0\xcf\x11\xe0":
        return "xls"

    # Neither format — capture a text preview for the log
    try:
        with open(path, encoding="utf-8", errors="replace") as fh:
            snippet = fh.read(500)
    except Exception:
        snippet = repr(header)

    log.error(f"Downloaded file is not a valid Excel file. Header bytes: {header!r}")
    log.error(f"File content preview:\n{snippet}")
    raise RuntimeError(
        "Download fehlgeschlagen – Mergermarket hat keine Excel-Datei geliefert.\n"
        f"Datei-Anfang: {snippet[:200]}"
    )


def _trigger_download(page, ctx, output_path: Path):
    """
    Mergermarket three-step download flow:

      1. Results page  →  click 'Download all'
      2. Download options page  →  click #btnUnformattedDownload
      3. Wait for 'Click here to view your report' link  →  click it
         (this final click triggers the actual file download)
    """
    ctx.wait_for_timeout(2_000)

    # ── Step 1: 'Download all' on the results page ───────────────────────────
    _dump_page_state(page, "04a_before_download_all")
    clicked = (
        _try_click(ctx, [
            "a:has-text('Download all')",
            "a:has-text('Download All')",
            "button:has-text('Download all')",
            "button:has-text('Download All')",
            "input[value='Download all']",
            "input[value='Download All']",
        ])
        or _js_click_by_text(ctx, "download all")
    )
    if not clicked:
        raise RuntimeError(
            "'Download all' element not found on the results page.\n"
            "Open C:\\Temp\\mm_debug_04a_before_download_all.png"
        )
    log.info("Clicked 'Download all'")
    page.wait_for_load_state("networkidle", timeout=15_000)

    # ── Step 2: click #btnUnformattedDownload ────────────────────────────────
    _dump_page_state(page, "04b_download_options")
    log.info("Clicking #btnUnformattedDownload …")
    clicked2 = _try_click(ctx, [
        "#btnUnformattedDownload",
        "input#btnUnformattedDownload",
        "button#btnUnformattedDownload",
    ])
    if not clicked2:
        raise RuntimeError(
            "#btnUnformattedDownload not found on the download options page.\n"
            "Check C:\\Temp\\mm_debug_04b_download_options.{png,json}"
        )
    log.info("Clicked #btnUnformattedDownload — waiting for 'Click here' link …")

    # ── Step 3: locate the 'Click here to view your report' link ─────────────
    #            Find it first, then click inside expect_download so the
    #            Playwright download handler is already registered when the
    #            browser starts the file transfer.
    log.info("Waiting for 'Click here to view your report' link (up to 60 s) …")
    link_el = None
    _LINK_SELECTORS = [
        "a:has-text('Click here to view your report')",
        "a:has-text('view your report')",
        "a:has-text('click here')",
    ]
    for attempt in range(30):  # 30 × 2 s = 60 s max
        for sel in _LINK_SELECTORS:
            try:
                link_el = ctx.wait_for_selector(sel, timeout=2_000)
                if link_el:
                    log.info(f"Found download link via {sel!r} (attempt {attempt + 1}/30)")
                    break
            except Exception:
                continue
        if link_el:
            break
        log.debug(f"Link not yet visible — waiting 2 s (attempt {attempt + 1}/30) …")
        ctx.wait_for_timeout(2_000)

    _dump_page_state(page, "04c_view_report_link")

    if not link_el:
        raise RuntimeError(
            "'Click here to view your report' link not found after clicking "
            "#btnUnformattedDownload.\n"
            f"Check {TEMP_DIR}\\mm_debug_04c_view_report_link.png"
        )

    log.info("Clicking download link inside expect_download context …")
    with page.expect_download(timeout=90_000) as dl_info:
        link_el.click()

    download = dl_info.value
    log.info(f"Download event captured — suggested filename: {download.suggested_filename!r}")
    download.save_as(str(output_path))
    fmt = _validate_excel_download(output_path)

    # Rename to the correct extension if the server delivered .xls (OLE2)
    # but the placeholder path has .xlsx
    if fmt != output_path.suffix.lstrip(".").lower():
        correct_path = output_path.with_suffix(f".{fmt}")
        correct_path.unlink(missing_ok=True)  # Windows rename fails if target exists
        output_path.rename(correct_path)
        output_path = correct_path
        log.info(f"Renamed to .{fmt}: {output_path}")

    log.info(f"Download saved and validated ({fmt}): {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# Bundeskartellamt scraping (Friday only)
# ---------------------------------------------------------------------------


def scrape_bundeskartellamt(page) -> list[dict]:
    """
    Load the BKA Laufende Verfahren page in its default view (first 15 entries)
    and return the first 15 data rows without any Aktenzeichen filtering.

    Columns are positional: 0=Datum, 1=Aktenzeichen, 2=Unternehmen,
    3=Produktbereich, 4=Abschluss.
    """
    log.info("BKA: loading Laufende Verfahren page …")
    page.goto(BKA_URL, wait_until="domcontentloaded", timeout=30_000)
    page.wait_for_timeout(1_500)
    _dump_page_state(page, "bka_01_initial")

    # Parse every <tr> with <td> cells — use the default page size (15 entries).
    all_rows: list[list[str]] = page.evaluate("""() =>
        Array.from(document.querySelectorAll('table tr'))
            .map(tr => Array.from(tr.querySelectorAll('td'))
                           .map(td => td.textContent.trim()))
            .filter(cells => cells.length >= 2)
    """)

    log.info(f"BKA: {len(all_rows)} data rows found; taking first 15")

    results: list[dict] = []
    for cells in all_rows[:15]:
        results.append({
            "datum":          cells[0] if len(cells) > 0 else "",
            "aktenzeichen":   " ".join(cells[1].split()) if len(cells) > 1 else "",
            "unternehmen":    cells[2] if len(cells) > 2 else "",
            "produktbereich": cells[3] if len(cells) > 3 else "",
            "abschluss":      cells[4] if len(cells) > 4 else "",
        })

    log.info(f"BKA: returning {len(results)} rows")
    return results


# ---------------------------------------------------------------------------
# Step 2 – Parse the Raw Excel Report
# ---------------------------------------------------------------------------


def parse_excel_report(xlsx_path: Path) -> list[str]:
    """
    Read the Mergermarket unformatted report using its actual column structure:

      Row 1 : empty
      Row 2 : headers — col E = 'Heading', col F = full article text
      Row 3+: data

    Returns a list of newline-joined strings, one per report.
    The first line is the heading (col E); the second line is the article body
    (col F).  Rows where the Heading cell is empty or missing are skipped.
    """
    log.info(f"Parsing Excel report: {xlsx_path}")
    suffix = xlsx_path.suffix.lower()

    entries: list[str] = []

    if suffix == ".xls":
        try:
            import xlrd
        except ImportError:
            show_error(
                "Mergermarket – Missing dependency",
                "xlrd is not installed.\nRun: pip install xlrd",
            )
            raise
        wb = xlrd.open_workbook(str(xlsx_path))
        ws = wb.sheet_by_index(0)
        # row 0 = empty, row 1 = headers → data starts at row index 2
        for row_idx in range(2, ws.nrows):
            heading = str(ws.cell_value(row_idx, 4)).strip()   # col E
            if not heading or heading.lower() == "nan":
                continue
            body_raw = ws.cell_value(row_idx, 5)               # col F
            body = str(body_raw).strip() if body_raw not in (None, "") else ""
            entries.append(f"{heading}\n{body}" if body else heading)
        wb.release_resources()
    else:
        try:
            import openpyxl
        except ImportError:
            show_error(
                "Mergermarket – Missing dependency",
                "openpyxl is not installed.\nRun: pip install openpyxl",
            )
            raise
        wb = openpyxl.load_workbook(str(xlsx_path), read_only=True, data_only=True)
        ws = wb.active
        # row 1 = empty, row 2 = headers → data starts at row 3 (1-based)
        for row_idx in range(3, (ws.max_row or 502) + 1):
            heading_val = ws.cell(row=row_idx, column=5).value  # col E
            if heading_val is None or str(heading_val).strip() == "":
                continue
            heading = str(heading_val).strip()
            body_val = ws.cell(row=row_idx, column=6).value     # col F
            body = str(body_val).strip() if body_val is not None else ""
            entries.append(f"{heading}\n{body}" if body else heading)
        wb.close()

    log.info(f"Found {len(entries)} report entries.")
    return entries


# ---------------------------------------------------------------------------
# Step 3 – Generate the Formatted Word Document
# ---------------------------------------------------------------------------


def _add_hyperlink_to_top(paragraph) -> None:
    """Append a '(Top)' hyperlink pointing to the _top bookmark."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), "Top")
    hyperlink.set(qn("w:history"), "1")

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = "(Top)"
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def _add_bookmark(paragraph, bookmark_id: int, name: str) -> None:
    """Wrap paragraph content in a named Word bookmark."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = paragraph._p
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, bm_start)
    p.append(bm_end)


def _add_toc_hyperlink(paragraph, text: str, anchor: str) -> None:
    """Insert a hyperlink into a paragraph pointing to an internal bookmark."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def _apply_heading_formatting(paragraph, font_name: str = "Aptos") -> None:
    """Set Heading 1 style then override with Aptos 12pt Bold Black."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    paragraph.style = "Heading 1"
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Ensure the font name is stored in the rFonts XML element
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)


def generate_word_document(
    entries: list[str],
    output_path: Path,
    run_date: date,
) -> Path:
    """
    Build the formatted .docx from a list of report entry strings.

    Structure: Top bookmark → numbered TOC → separator/Heading/body/(Top) per report.
    """
    try:
        from docx import Document
    except ImportError:
        show_error("Mergermarket – Missing dependency", "python-docx is not installed.\nRun: pip install python-docx")
        raise

    log.info(f"Generating Word document: {output_path}")

    doc = Document()
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # Parse entries into (heading, body_lines) pairs
    parsed: list[tuple[str, list[str]]] = []
    for entry in entries:
        lines = [ln.strip() for ln in entry.splitlines() if ln.strip()]
        if not lines:
            continue
        parsed.append((lines[0], [ln for ln in lines[1:] if ln.strip() != "(Top)"]))

    if not parsed:
        log.warning("No report entries — the Word document will be empty.")
        show_error(
            "Mergermarket – No Reports",
            "No intelligence reports were found for today.\nCheck the downloaded Excel file.",
        )

    # ── Top bookmark (before TOC so (Top) links land at document start) ──────
    top_para = doc.add_paragraph()
    top_para.style = "Normal"
    _add_bookmark(top_para, bookmark_id=0, name="Top")

    # ── Numbered TOC: one hyperlink per report → _Toc_{i} ────────────────────
    for i, (heading, _) in enumerate(parsed, 1):
        toc_para = doc.add_paragraph()
        toc_para.style = "Normal"
        _add_toc_hyperlink(toc_para, f"{i}. {heading}", anchor=f"_Toc_{i}")

    doc.add_paragraph().style = "Normal"  # blank line after TOC

    # ── Reports ───────────────────────────────────────────────────────────────
    for i, (heading, body_lines) in enumerate(parsed, 1):
        doc.add_paragraph(SEPARATOR).style = "Normal"

        # Heading 1 in Aptos 12pt Bold with bookmark _Toc_{i}
        heading_para = doc.add_paragraph()
        heading_para.add_run(f"{i}. {heading}")
        _apply_heading_formatting(heading_para, font_name="Aptos")
        _add_bookmark(heading_para, bookmark_id=i, name=f"_Toc_{i}")

        for line in body_lines:
            doc.add_paragraph(line).style = "Normal"

        top_link_para = doc.add_paragraph()
        top_link_para.style = "Normal"
        _add_hyperlink_to_top(top_link_para)

    doc.add_paragraph(SEPARATOR).style = "Normal"

    doc.save(str(output_path))
    log.info(f"Word document saved: {output_path}  ({len(parsed)} reports)")

    _refresh_toc(output_path)

    return output_path


def _refresh_toc(docx_path: Path) -> None:
    """Open the saved .docx in Word via COM, update the TOC field, and save."""
    try:
        import win32com.client
    except ImportError:
        log.warning("pywin32 not available — TOC placeholder will remain; update manually.")
        return
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.TablesOfContents(1).Update()
        doc.Save()
        doc.Close()
        if word.Documents.Count == 0:
            word.Quit()
        log.info("TOC updated and document re-saved.")
    except Exception as exc:
        log.warning(f"TOC auto-update failed ({exc}) — open the document and press Ctrl+A → F9.")


# ---------------------------------------------------------------------------
# Step 4 – Compose and Display the Outlook Email
# ---------------------------------------------------------------------------


def compose_outlook_email(
    word_doc_path: Path,
    run_date: date,
    *,
    bka_data: list[dict] | None = None,
    is_friday: bool = False,
    auto_send: bool = False,
) -> None:
    """
    Create an Outlook MailItem with intro text, the Word document content
    pasted directly into the email body, and a closing signature.
    """
    try:
        import win32com.client
    except ImportError:
        show_error(
            "Mergermarket – Missing dependency",
            "pywin32 is not installed.\nRun: pip install pywin32",
        )
        raise

    if is_friday:
        subject = (
            f"Laufende Verfahren Bundeskartellamt und Mergermarket Newsletter"
            f" - {run_date.strftime('%d.%m.%Y')}"
        )
    else:
        subject = f"Mergermarket Newsletter - {run_date.strftime('%d.%m.%Y')}"

    import subprocess as _sp
    import glob as _glob

    try:
        tasklist_out = _sp.run(
            ["tasklist", "/NH"],
            capture_output=True, text=True, timeout=10,
        ).stdout.upper()
        outlook_running = "OLK.EXE" in tasklist_out or "OUTLOOK.EXE" in tasklist_out
    except Exception:
        outlook_running = False

    if not outlook_running:
        olk_candidates = _glob.glob(
            str(Path.home() / "AppData" / "Local" / "Microsoft" / "WindowsApps" / "olk.exe")
        )
        classic_candidates = (
            _glob.glob(r"C:\Program Files\Microsoft Office\root\Office*\OUTLOOK.EXE")
            + _glob.glob(r"C:\Program Files (x86)\Microsoft Office\root\Office*\OUTLOOK.EXE")
        )
        outlook_exe = (olk_candidates or classic_candidates or ["outlook.exe"])[0]
        log.info(f"Outlook not running — starting: {outlook_exe}")
        _sp.Popen([outlook_exe])
    else:
        log.info("Outlook already running — connecting via COM …")

    outlook = None
    for attempt in range(1, 31):
        try:
            outlook = win32com.client.Dispatch("Outlook.Application")
            log.info(f"Outlook COM connection established (attempt {attempt}/30)")
            break
        except Exception:
            log.info(f"Waiting for Outlook... (attempt {attempt}/30)")
            time.sleep(5)

    if outlook is None:
        msg = "Outlook did not become ready within 2.5 minutes."
        show_error("Mergermarket – Outlook Error", msg)
        raise RuntimeError(msg)

    mail = outlook.CreateItem(0)  # 0 = olMailItem
    mail.Subject = subject

    for addr in [EMAIL_RECIPIENT, "sonke.debuhr@casecassiopea.com"]:
        r = mail.Recipients.Add(addr)
        r.Resolve()

    # Open the Word document via COM to copy its content later
    log.info("Opening Word document via COM …")
    try:
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        source_doc = word_app.Documents.Open(str(word_doc_path.resolve()))
    except Exception as exc:
        show_error("Mergermarket – Word Error", f"Could not open the Word document:\n{exc}")
        raise

    mail.Display()

    inspector = mail.GetInspector
    mail_doc = inspector.WordEditor
    word_selection = mail_doc.Application.Selection

    # Save auto-signature before clearing body (will be re-appended at end)
    sig_text = mail_doc.Range().Text.strip()
    mail_doc.Range().Delete()

    try:
        display_name = outlook.Session.CurrentUser.Name
        first_name = display_name.split()[0] if display_name else ""
    except Exception:
        first_name = ""
    if len(first_name) <= 1:
        win_user = os.environ.get("USERNAME", "")
        first_name = win_user.split(".")[0].capitalize() if win_user else first_name

    def _body_font() -> None:
        word_selection.Font.Name = "Aptos"
        word_selection.Font.Size = 12
        word_selection.Font.Bold = False

    def _bold(text: str) -> None:
        word_selection.Font.Bold = True
        word_selection.TypeText(text)
        word_selection.Font.Bold = False

    word_selection.HomeKey(Unit=6)  # wdStory = 6
    _body_font()

    if is_friday and bka_data:
        # ── Friday intro ──────────────────────────────────────────────────
        for line in FRIDAY_INTRO.splitlines():
            word_selection.TypeText(line)
            word_selection.TypeParagraph()
        word_selection.TypeParagraph()

        _bold("Bundeskartellamt:")
        word_selection.TypeParagraph()
        word_selection.TypeParagraph()

        _BKA_HEADERS = ["Datum", "Aktenzeichen", "Unternehmen",
                        "Produktbereich", "Abschluss"]
        _BKA_KEYS    = ["datum", "aktenzeichen", "unternehmen",
                        "produktbereich", "abschluss"]
        tbl = mail_doc.Tables.Add(
            Range=word_selection.Range,
            NumRows=1 + len(bka_data),
            NumColumns=len(_BKA_HEADERS),
        )

        _DARK_BLUE = 0x68 * 65536
        _WHITE     = 0xFF + 0xFF * 256 + 0xFF * 65536
        for j, h in enumerate(_BKA_HEADERS, 1):
            cell = tbl.Cell(1, j)
            cell.Range.Text = h
            cell.Range.Font.Bold = True
            cell.Range.Font.Color = _WHITE
            cell.Shading.BackgroundPatternColor = _DARK_BLUE

        for i, row in enumerate(bka_data, 2):
            for j, key in enumerate(_BKA_KEYS, 1):
                tbl.Cell(i, j).Range.Text = str(row.get(key, ""))

        _col_width_pt = mail_doc.Application.PixelsToPoints(550, False)
        tbl.Columns(3).Width = _col_width_pt
        tbl.Columns(4).Width = _col_width_pt
        tbl.Borders.OutsideLineStyle = 1
        tbl.Borders.OutsideColor = 0
        for _row_idx in range(1, tbl.Rows.Count + 1):
            for _col_idx in range(1, tbl.Columns.Count + 1):
                _cell = tbl.Cell(_row_idx, _col_idx)
                _cell.Borders(3).LineStyle = 1  # bottom border per cell
                _cell.Borders(3).Color = 0

        word_selection.EndKey(Unit=6)
        word_selection.TypeParagraph()
        word_selection.TypeParagraph()
        _body_font()

        _bold("Mergermarket:")
        word_selection.TypeParagraph()
        _body_font()

    else:
        # ── Normal day intro ──────────────────────────────────────────────
        for line in EMAIL_INTRO.splitlines():
            if line == "Mergermarket:":
                _bold(line)
            else:
                word_selection.TypeText(line)
            word_selection.TypeParagraph()
        _body_font()

    # ── Paste Word document content ───────────────────────────────────────
    source_doc.Content.Copy()
    word_selection.EndKey(Unit=6)
    word_selection.Paste()

    # Clear the clipboard so Word doesn't ask to keep large clipboard contents
    import subprocess as _sub
    _sub.run(
        ["powershell", "-command", "Set-Clipboard -Value $null"],
        capture_output=True,
    )
    source_doc.Close(SaveChanges=False)
    if word_app.Documents.Count == 0:
        word_app.Quit()

    # ── Closing signature ─────────────────────────────────────────────────
    word_selection.EndKey(Unit=6)
    word_selection.TypeParagraph()
    word_selection.TypeText("Beste Grüße")
    word_selection.TypeParagraph()
    word_selection.TypeText(first_name)

    # Re-append Outlook auto-signature at the very end
    if sig_text:
        word_selection.TypeParagraph()
        word_selection.TypeParagraph()
        for line in sig_text.splitlines():
            word_selection.TypeText(line)
            word_selection.TypeParagraph()

    if auto_send:
        mail.Send()
        log.info("Email sent automatically.")
    else:
        log.info(f"Email displayed for manual review (signed as {first_name!r}).")


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------


def run(
    run_date: date,
    dry_run_xlsx: Optional[Path] = None,
    headless: bool = False,
    force_friday: bool = False,
    auto_send: bool = False,
) -> None:
    today_str = run_date.strftime("%Y%m%d")
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    is_friday = force_friday or (run_date.weekday() == 4)
    if is_friday:
        log.info("Friday mode active — BKA scraping enabled.")

    raw_xlsx = TEMP_DIR / f"mergermarket_raw_{today_str}.xlsx"        # temp
    output_docx = OUTPUT_DIR / f"mergermarket_report_{today_str}.docx"  # Downloads

    try:
        # ── Step 1: Download ─────────────────────────────────────────────────
        if dry_run_xlsx:
            raw_xlsx = dry_run_xlsx
            bka_data: list[dict] = []
            log.info(f"[DRY-RUN] Skipping browser download; using: {raw_xlsx}")
        else:
            raw_xlsx, bka_data = download_mergermarket_report(
                run_date, raw_xlsx, headless=headless, is_friday=is_friday
            )

        # ── Step 2: Parse ────────────────────────────────────────────────────
        entries = parse_excel_report(raw_xlsx)
        if not entries:
            show_error(
                "Mergermarket – No Data",
                f"No report entries were found in:\n{raw_xlsx}",
            )
            return

        # ── Step 3: Format Word document ─────────────────────────────────────
        generate_word_document(entries, output_docx, run_date)

        # ── Step 4: Compose email ────────────────────────────────────────────
        compose_outlook_email(
            output_docx, run_date, bka_data=bka_data,
            is_friday=is_friday, auto_send=auto_send,
        )

    except Exception as exc:
        log.exception("Fatal error in Mergermarket newsletter automation")
        show_error(
            "Mergermarket – Fatal Error",
            f"The automation failed:\n{exc}\n\nFull log: {LOG_FILE}",
        )
        sys.exit(1)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Automate the Mergermarket daily newsletter workflow.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "--dry-run",
        metavar="XLSX_PATH",
        help="Skip browser download and use an existing Excel file instead.",
    )
    parser.add_argument(
        "--date",
        metavar="YYYY-MM-DD",
        help="Override today's date (e.g. for a Monday run covering the weekend).",
    )
    parser.add_argument(
        "--headless",
        action="store_true",
        help="Run the browser in headless mode (no visible window).",
    )
    parser.add_argument(
        "--friday",
        action="store_true",
        help="Force Friday mode: scrape Bundeskartellamt and use the extended email format.",
    )
    parser.add_argument(
        "--send",
        action="store_true",
        help="Send the email automatically via mail.Send(). Without this flag the email is only displayed for manual review.",
    )
    args = parser.parse_args()

    run_date = get_run_date(args.date)
    log.info(f"=== Mergermarket Newsletter  {run_date}  ({'Monday' if run_date.weekday() == 0 else run_date.strftime('%A')}) ===")

    if args.dry_run:
        run(run_date, dry_run_xlsx=Path(args.dry_run), headless=args.headless,
            force_friday=args.friday, auto_send=args.send)
    else:
        run(run_date, headless=args.headless, force_friday=args.friday, auto_send=args.send)


if __name__ == "__main__":
    main()
